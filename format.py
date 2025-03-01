import streamlit as st
import openai
import docx
import io
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_formatting_from_docx(doc):
    """Extract formatting details from a reference Word document"""
    formatting = {
        "paragraphs": [],
        "headings": [],
        "tables": [],
        "overall_style": {}
    }
    
    # Extract paragraph formatting
    for para in doc.paragraphs:
        if para.text.strip():  # Skip empty paragraphs
            style_info = {
                "text": para.text,
                "alignment": para.alignment,
                "style_name": para.style.name if para.style else "Normal",
                "font_properties": []
            }
            
            # Extract run properties (font, size, bold, italic, etc.)
            for run in para.runs:
                if run.text.strip():
                    run_props = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name if run.font else None,
                        "font_size": run.font.size.pt if run.font.size else None,
                        "color": run.font.color.rgb if run.font.color and run.font.color.rgb else None
                    }
                    style_info["font_properties"].append(run_props)
            
            # Categorize as heading or paragraph
            if para.style and "Heading" in para.style.name:
                formatting["headings"].append(style_info)
            else:
                formatting["paragraphs"].append(style_info)
    
    # Extract table formatting - fixed to avoid the get_or_add_shd error
    for table in doc.tables:
        table_info = {
            "rows": len(table.rows),
            "cols": len(table.columns),
            "cell_styles": []
        }
        
        for row in table.rows:
            for cell in row.cells:
                cell_info = {
                    "text": cell.text,
                    "vertical_alignment": cell.vertical_alignment
                }
                # We'll skip the shading extraction that was causing the error
                table_info["cell_styles"].append(cell_info)
        
        formatting["tables"].append(table_info)
    
    # Extract document-wide properties
    formatting["overall_style"] = {
        "default_font": doc.styles["Normal"].font.name if "Normal" in doc.styles else None,
        "default_font_size": doc.styles["Normal"].font.size.pt if "Normal" in doc.styles and doc.styles["Normal"].font.size else None,
        "margins": {
            "top": doc.sections[0].top_margin.inches if doc.sections else None,
            "bottom": doc.sections[0].bottom_margin.inches if doc.sections else None,
            "left": doc.sections[0].left_margin.inches if doc.sections else None,
            "right": doc.sections[0].right_margin.inches if doc.sections else None
        }
    }
    
    return formatting

def apply_formatting_to_docx(target_doc, formatting_info, openai_client):
    """
    Apply formatting from reference document to target document using OpenAI.
    This function will use OpenAI to understand the document structure and 
    then apply the formatting from the reference document.
    """
    # Extract content from target document
    content = []
    for para in target_doc.paragraphs:
        if para.text.strip():
            content.append({"type": "paragraph", "text": para.text})
    
    # Create a new document that will have the formatting applied
    new_doc = docx.Document()
    
    # Set document-wide properties
    overall_style = formatting_info["overall_style"]
    if overall_style.get("margins"):
        for section in new_doc.sections:
            if overall_style["margins"].get("top"):
                section.top_margin = Inches(overall_style["margins"]["top"])
            if overall_style["margins"].get("bottom"):
                section.bottom_margin = Inches(overall_style["margins"]["bottom"])
            if overall_style["margins"].get("left"):
                section.left_margin = Inches(overall_style["margins"]["left"])
            if overall_style["margins"].get("right"):
                section.right_margin = Inches(overall_style["margins"]["right"])
    
    # Use OpenAI to identify paragraphs vs headings in the target document
    prompt = """
    Analyze the following document content and classify each paragraph as either 'heading' or 'body':
    
    """
    for item in content:
        prompt += f"\nText: {item['text']}\n"
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a document formatting assistant. Classify each paragraph as 'heading' or 'body' based on its content and context."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500
        )
        
        classifications = response.choices[0].message.content
        
        # Parse OpenAI's response to get classifications
        lines = classifications.lower().split("\n")
        classified_paragraphs = []
        
        # Improved parsing to handle various response formats
        for i, item in enumerate(content):
            # Try to find a classification for this paragraph
            classification = None
            for line in lines:
                if item["text"][:20].lower() in line.lower():
                    if "heading" in line.lower():
                        classification = "heading"
                    else:
                        classification = "body"
                    break
            
            # If no specific match found, use index-based approach as fallback
            if classification is None and i < len(lines):
                if "heading" in lines[i].lower():
                    classification = "heading"
                else:
                    classification = "body"
            
            # Default to body if still no classification
            if classification is None:
                classification = "body"
                
            classified_paragraphs.append({"type": classification, "text": item["text"]})
        
        # Apply formatting based on classifications
        for item in classified_paragraphs:
            if item["type"] == "heading":
                # Apply heading formatting
                if formatting_info["headings"]:
                    ref_heading = formatting_info["headings"][0]
                    p = new_doc.add_paragraph()
                    run = p.add_run(item["text"])
                    
                    if ref_heading.get("alignment") is not None:
                        p.alignment = ref_heading["alignment"]
                    
                    # Try to find the heading style
                    if ref_heading.get("style_name") and ref_heading["style_name"] in new_doc.styles:
                        p.style = ref_heading["style_name"]
                    else:
                        # Default to Heading 1 if the specific style doesn't exist
                        try:
                            p.style = "Heading 1"
                        except:
                            # If Heading 1 style doesn't exist, just make it bold and larger
                            run.bold = True
                            run.font.size = Pt(16)
                    
                    # Apply character formatting from reference
                    if ref_heading["font_properties"]:
                        font_props = ref_heading["font_properties"][0]
                        if font_props.get("bold") is not None:
                            run.bold = font_props["bold"]
                        if font_props.get("italic") is not None:
                            run.italic = font_props["italic"]
                        if font_props.get("font_name"):
                            run.font.name = font_props["font_name"]
                        if font_props.get("font_size"):
                            run.font.size = Pt(font_props["font_size"])
                        if font_props.get("color"):
                            try:
                                run.font.color.rgb = font_props["color"]
                            except:
                                # Handle case where color can't be set
                                pass
                else:
                    # Default heading formatting if no reference
                    p = new_doc.add_paragraph()
                    run = p.add_run(item["text"])
                    run.bold = True
                    try:
                        p.style = "Heading 1"
                    except:
                        run.font.size = Pt(16)
            else:
                # Apply body paragraph formatting
                if formatting_info["paragraphs"]:
                    ref_para = formatting_info["paragraphs"][0]
                    p = new_doc.add_paragraph()
                    run = p.add_run(item["text"])
                    
                    if ref_para.get("alignment") is not None:
                        p.alignment = ref_para["alignment"]
                    
                    # Apply character formatting from reference
                    if ref_para["font_properties"]:
                        font_props = ref_para["font_properties"][0]
                        if font_props.get("bold") is not None:
                            run.bold = font_props["bold"]
                        if font_props.get("italic") is not None:
                            run.italic = font_props["italic"]
                        if font_props.get("font_name"):
                            run.font.name = font_props["font_name"]
                        if font_props.get("font_size"):
                            run.font.size = Pt(font_props["font_size"])
                        if font_props.get("color"):
                            try:
                                run.font.color.rgb = font_props["color"]
                            except:
                                # Handle case where color can't be set
                                pass
                else:
                    # Default paragraph formatting if no reference
                    p = new_doc.add_paragraph(item["text"])
                    
        return new_doc
    
    except Exception as e:
        st.error(f"Error using OpenAI API: {str(e)}")
        return None

def main():
    st.title("Document Formatter with OpenAI")
    st.write("Upload a reference document and a target document to apply the same formatting.")
    
    api_key = st.text_input("Enter your OpenAI API key", type="password")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Reference Document (with formatting)")
        reference_file = st.file_uploader("Upload reference DOCX file", type=["docx"])
    
    with col2:
        st.subheader("Target Document (to be formatted)")
        target_file = st.file_uploader("Upload target DOCX file", type=["docx"])
    
    if api_key and reference_file and target_file:
        try:
            # Set up OpenAI client
            client = openai.OpenAI(api_key=api_key)
            
            # Load documents
            reference_doc = docx.Document(reference_file)
            target_doc = docx.Document(target_file)
            
            if st.button("Format Document"):
                with st.spinner("Extracting formatting from reference document..."):
                    formatting_info = extract_formatting_from_docx(reference_doc)
                
                with st.spinner("Applying formatting to target document..."):
                    formatted_doc = apply_formatting_to_docx(target_doc, formatting_info, client)
                
                if formatted_doc:
                    st.success("Document formatted successfully!")
                    
                    # Save the document to a BytesIO object for download
                    doc_bytes = io.BytesIO()
                    formatted_doc.save(doc_bytes)
                    doc_bytes.seek(0)
                    
                    # Create download button
                    st.download_button(
                        label="Download Formatted Document",
                        data=doc_bytes,
                        file_name="formatted_document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
    else:
        st.info("Please provide your OpenAI API key and upload both documents.")

if __name__ == "__main__":
    main()